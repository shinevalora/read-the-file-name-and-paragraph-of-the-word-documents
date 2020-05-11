# _*_coding:utf-8_*_
#  作者    : shinevalora
#  创建时间: 2020/5/8  11:08

import os
import re
import logging
from datetime import datetime
from multiprocessing import Process

from docx import Document

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)-8s: %(message)s")


def handle_file(file):
    '''
    单个word文件处理
    '''
    if os.path.isfile(file):
        # 分割出文件路径和文件后缀名
        pre, ext = os.path.splitext(file)
        # 过滤出后缀名为docx的word文件和去除以 ~$ 开头的docx文件
        if ext == ".docx" and "~$" not in file:
            # 使用正则替换掉文字字符串，最后只剩下数字(日期信息)
            pattern = re.sub(r'\D', '', pre)
            # 读取word文档
            word_file = Document(file)

            with open('experiment_title.txt', 'a+', encoding="utf-8") as f:
                # 分割出文件路径
                data = [_date for _date in file.split("\\")]
                for _data in data:
                    if ".docx" in _data:
                        # 读取word文件中文件名称的前八位内容和第三段内容和并保存至 txt文件中
                        f.write(_data[:8] + "    " + word_file.paragraphs[3].text + "\n")
                        # 日志显示
                        logging.info(_data[:8] + "    " + word_file.paragraphs[3].text)


def handle_dir(path):
    '''
    多层级文件夹遍历
    :path: 多层级文件夹
    :return: 返回单个word文件
    '''

    # 路径不存在直接退
    if not os.path.exists(path):
        return
    # 如果路径是文件夹，那么获取路径文件夹下面所有的文件和文件夹，然后遍历处理
    if os.path.isdir(path):
        for child_dir_or_file in os.listdir(path):
            child_path = os.path.join(path, child_dir_or_file)
            if os.path.isfile(child_path):
                handle_file(child_path)
            else:
                # 递归
                handle_dir(child_path)
    else:
        if os.path.splitext(path)[1] == ".docx" and "~$" not in path:
            handle_file(path)


def folders_test(num):
    '''
    测试用文件夹
    :num: 创建多少个文件夹
    :return: 会生成一个  test_folders 文件夹
    '''

    for i in range(1, num + 1):
        folders = f"test_folders/test_folders{str(i).zfill(2)}"
        os.makedirs(folders)
        # logging.info(folders,type(folders))
        _ = 1
        for j in range(1, num+1):
            # 创建Document对象
            doc = Document()
            _date = datetime.now().strftime("%Y%m%d")
            doc.add_paragraph(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')} {_} 第一段内容")
            doc.add_paragraph(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')} {_} 第二段内容")
            doc.add_paragraph(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')} {_} 第三段内容")
            doc.add_paragraph(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')} {_} 第四段内容")

            _ += 1

            doc.save(f'{folders}/{datetime.now().strftime("%Y%m%d")}_{j}.docx')

            logging.info(f"正在生成和保存  {folders}/{_date}_{j}.docx")


if __name__ == "__main__":
    # 测试用文件夹，如自己有文件夹需读取，无需对此操作
    p1 = Process(target=folders_test, args=(12,))
    p1.start()
    p1.join()

    p = Process(target=handle_dir, args=(r"test_folders",))
    p.start()
    p.join()
